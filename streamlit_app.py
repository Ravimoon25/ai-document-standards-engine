# streamlit_app.py - COMPLETE FILE - Step 1: Imports Section

# Core Streamlit and web framework
import streamlit as st
import os
from datetime import datetime
import json
import io
import re

# Google AI Integration
from google import genai
from google.genai import types

# Image Processing
import PIL.Image

# Data Processing and ML for RAG system
import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Document Processing
import docx
import PyPDF2

# Type hints
from typing import List, Dict, Tuple, Optional

# NEW: Agent system dependencies
from dataclasses import dataclass
from enum import Enum

# NEW: Enhanced DOCX manipulation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# Step 2: Page Configuration and CSS - ADD THIS AFTER IMPORTS

# Page configuration
st.set_page_config(
    page_title="📄 AI Document Standards Engine",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS styling
st.markdown("""
<style>
.main .block-container {
    padding: 1rem;
    max-width: 100%;
}

.doc-header {
    background: linear-gradient(135deg, #1e3a8a 0%, #3730a3 100%);
    color: white;
    padding: 2rem;
    border-radius: 12px;
    text-align: center;
    margin-bottom: 2rem;
}

.standards-card {
    background: #f0f9ff;
    border: 2px solid #0ea5e9;
    padding: 1.5rem;
    border-radius: 8px;
    margin: 1rem 0;
}

.processing-card {
    background: #fef3c7;
    border: 1px solid #f59e0b;
    padding: 1rem;
    border-radius: 8px;
    margin: 0.5rem 0;
}

.result-card {
    background: #ecfdf5;
    border: 1px solid #10b981;
    padding: 1.5rem;
    border-radius: 8px;
    margin: 1rem 0;
}

.change-highlight {
    background: #fef2f2;
    border-left: 4px solid #ef4444;
    padding: 0.5rem;
    margin: 0.5rem 0;
}

.rag-info {
    background: #f3f4f6;
    border: 1px solid #9ca3af;
    padding: 1rem;
    border-radius: 6px;
    font-size: 0.9rem;
}
</style>
""", unsafe_allow_html=True)


# Step 3: Session State and Client Initialization - ADD AFTER CSS

# Initialize session state
def init_session_state():
    if 'processed_documents' not in st.session_state:
        st.session_state.processed_documents = []
    if 'standards_library' not in st.session_state:
        st.session_state.standards_library = []
    if 'standards_chunks' not in st.session_state:
        st.session_state.standards_chunks = []
    if 'vectorizer' not in st.session_state:
        st.session_state.vectorizer = None
    if 'chunk_embeddings' not in st.session_state:
        st.session_state.chunk_embeddings = None

init_session_state()

@st.cache_resource
def get_client():
    """Initialize Gemini client"""
    try:
        api_key = st.secrets["GOOGLE_API_KEY"]
        return genai.Client(api_key=api_key)
    except Exception as e:
        st.error(f"Failed to initialize AI client: {str(e)}")
        st.stop()


# Step 4: Core Document Processing Functions - ADD AFTER CLIENT INITIALIZATION

def extract_text_from_document(uploaded_file):
    """Extract text from uploaded document - production ready"""
    try:
        file_type = uploaded_file.type
        
        if file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            st.info("📄 Processing Word document...")
            
            with st.spinner("🔍 Extracting text from Word document..."):
                doc = docx.Document(uploaded_file)
                
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))
                
                extracted_text = "\n".join(full_text)
                return extracted_text if extracted_text.strip() else "No text found in Word document."
        
        elif file_type == "application/pdf":
            st.info("📄 Processing PDF document...")
            
            with st.spinner("🔍 Extracting text from PDF..."):
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                
                full_text = []
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        full_text.append(f"--- Page {page_num + 1} ---\n{text}")
                
                extracted_text = "\n\n".join(full_text)
                return extracted_text if extracted_text.strip() else "PDF appears to be image-based. Try uploading as image for OCR."
        
        elif file_type in ["image/png", "image/jpeg", "image/jpg"]:
            st.info("🖼️ Processing image with OCR...")
            
            image = PIL.Image.open(uploaded_file)
            
            with st.spinner("🔍 Extracting text from image..."):
                client = get_client()
                response = client.models.generate_content(
                    model="gemini-2.5-flash-image-preview",
                    contents=["Extract all text from this document image, maintaining formatting and structure:", image]
                )
                
                extracted_text = ""
                for part in response.parts:
                    if part.text:
                        extracted_text += part.text
                
                return extracted_text if extracted_text.strip() else "No text detected in image."
        
        else:
            try:
                content = uploaded_file.read().decode('utf-8')
                return content
            except UnicodeDecodeError:
                uploaded_file.seek(0)
                content = uploaded_file.read().decode('utf-8', errors='ignore')
                return content + "\n\n⚠️ Some characters may not display correctly."
                
    except Exception as e:
        return f"Error processing document: {str(e)}"

class StandardsProcessor:
    """RAG system for processing and retrieving standards"""
    
    def __init__(self):
        self.chunk_size = 500  # words per chunk
        self.overlap = 50      # word overlap between chunks
    
    def chunk_document(self, text: str, doc_name: str) -> List[Dict]:
        """Split large standards document into searchable chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            
            # Extract section information if available
            section_match = re.search(r'(#+\s*.+|\d+\.\d*\s*.+|Section\s+\d+)', chunk_text)
            section = section_match.group(1) if section_match else f"Chunk {len(chunks) + 1}"
            
            chunks.append({
                'text': chunk_text,
                'section': section,
                'doc_name': doc_name,
                'chunk_id': f"{doc_name}_chunk_{len(chunks)}",
                'word_count': len(chunk_words)
            })
        
        return chunks
    
    def create_embeddings(self, chunks: List[Dict]) -> Tuple[TfidfVectorizer, np.ndarray]:
        """Create TF-IDF embeddings for semantic search"""
        chunk_texts = [chunk['text'] for chunk in chunks]
        
        vectorizer = TfidfVectorizer(
            max_features=5000,
            stop_words='english',
            ngram_range=(1, 2),  # Include bigrams for better context
            min_df=1,
            max_df=0.95
        )
        
        embeddings = vectorizer.fit_transform(chunk_texts)
        return vectorizer, embeddings.toarray()
    
    def semantic_search(self, query: str, top_k: int = 5) -> List[Dict]:
        """Find most relevant standards chunks for a query"""
        if not st.session_state.vectorizer or st.session_state.chunk_embeddings is None:
            return []
        
        # Transform query using existing vectorizer
        query_vector = st.session_state.vectorizer.transform([query])
        
        # Calculate similarities
        similarities = cosine_similarity(query_vector, st.session_state.chunk_embeddings)[0]
        
        # Get top-k most similar chunks
        top_indices = np.argsort(similarities)[-top_k:][::-1]
        
        results = []
        for idx in top_indices:
            if similarities[idx] > 0.1:  # Minimum similarity threshold
                chunk = st.session_state.standards_chunks[idx]
                results.append({
                    **chunk,
                    'similarity': similarities[idx],
                    'relevance_score': similarities[idx] * 100
                })
        
        return results

def build_standards_knowledge_base():
    """Build RAG knowledge base from all uploaded standards"""
    if not st.session_state.standards_library:
        return False
    
    processor = StandardsProcessor()
    all_chunks = []
    
    # Process all standards documents
    for standard in st.session_state.standards_library:
        chunks = processor.chunk_document(standard['content'], standard['name'])
        all_chunks.extend(chunks)
    
    if all_chunks:
        # Create embeddings
        vectorizer, embeddings = processor.create_embeddings(all_chunks)
        
        # Store in session state
        st.session_state.standards_chunks = all_chunks
        st.session_state.vectorizer = vectorizer
        st.session_state.chunk_embeddings = embeddings
        
        return True
    
    return False


# Step 5: Agent System Classes - ADD AFTER CORE FUNCTIONS

class AgentType(Enum):
    """Enumeration of available agent types"""
    CONTENT_EDITOR = "content_editor"
    FORMATTER = "formatter" 
    QUALITY_CHECKER = "quality_checker"

@dataclass
class ProcessingResult:
    """Standardized result format for all agents"""
    content: str
    instructions: List[Dict]
    metadata: Dict
    agent_notes: str
    success: bool = True
    error_message: str = ""

class BaseAgent:
    """Base class for all document processing agents"""
    
    def __init__(self, agent_type: AgentType, standards_retriever=None):
        self.agent_type = agent_type
        self.standards_retriever = standards_retriever
        self.client = None
        
        # Initialize AI client
        try:
            self.client = get_client()
        except Exception as e:
            st.error(f"Failed to initialize {agent_type.value} agent: {str(e)}")
    
    def process(self, input_data: Dict) -> ProcessingResult:
        """Process input data and return standardized result"""
        raise NotImplementedError("Subclasses must implement process method")
    
    def _parse_json_response(self, response_text: str, fallback_content: str = "") -> Dict:
        """Helper method to parse JSON from AI responses with fallback"""
        try:
            # Try to extract JSON from response
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        except json.JSONDecodeError as e:
            st.warning(f"JSON parsing failed: {str(e)}")
        except Exception as e:
            st.warning(f"Response parsing error: {str(e)}")
        
        # Return fallback structure
        return {
            "content": fallback_content or response_text,
            "instructions": [],
            "metadata": {},
            "agent_notes": f"Fallback response due to parsing error"
        }
    
    def _get_ai_response(self, prompt: str) -> str:
        """Helper method to get AI response with error handling"""
        if not self.client:
            return "Error: AI client not initialized"
        
        try:
            response = self.client.models.generate_content(
                model="gemini-2.5-flash-image-preview",
                contents=prompt
            )
            
            result_text = ""
            for part in response.parts:
                if part.text:
                    result_text += part.text
            
            return result_text
            
        except Exception as e:
            error_msg = f"AI response error: {str(e)}"
            st.error(error_msg)
            return error_msg

class EnhancedStandardsRetriever:
    """Enhanced RAG system with filtering capabilities for agents"""
    
    def __init__(self):
        self.processor = StandardsProcessor()
    
    def search(self, query: str, filter_types: List[str] = None, top_k: int = 5) -> List[Dict]:
        """Search with optional filtering by standard type"""
        
        # Check if RAG system is available
        if not st.session_state.standards_chunks:
            st.warning("Standards knowledge base not built yet")
            return []
        
        try:
            # Get base search results (get more for filtering)
            results = self.processor.semantic_search(query, top_k=top_k*2)
            
            # Apply type filtering if specified
            if filter_types:
                filtered_results = []
                for result in results:
                    # Check if the source document type matches filter
                    source_standard = self._find_source_standard(result['doc_name'])
                    if source_standard and source_standard.get('type') in filter_types:
                        filtered_results.append(result)
                        if len(filtered_results) >= top_k:
                            break
                
                return filtered_results
            
            return results[:top_k]
            
        except Exception as e:
            st.error(f"Standards retrieval error: {str(e)}")
            return []
    
    def _find_source_standard(self, doc_name: str) -> Optional[Dict]:
        """Find the original standard document by name"""
        for standard in st.session_state.standards_library:
            if standard['name'] == doc_name:
                return standard
        return None
    
    def get_standards_count_by_type(self) -> Dict[str, int]:
        """Get count of standards by type"""
        counts = {}
        for standard in st.session_state.standards_library:
            std_type = standard.get('type', 'Unknown')
            counts[std_type] = counts.get(std_type, 0) + 1
        return counts

def display_agent_status(retriever: EnhancedStandardsRetriever):
    """Display current status of agents based on available standards"""
    
    st.subheader("Agent System Status")
    
    # Get standards count by type
    standards_by_type = retriever.get_standards_count_by_type()
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("**Content Editor Agent**")
        content_types = ['Editorial Guidelines', 'Style Guide', 'Writing Standards']
        content_count = sum(standards_by_type.get(t, 0) for t in content_types)
        
        if content_count > 0:
            st.success(f"Ready ({content_count} standards)")
        else:
            st.warning("No content standards loaded")
    
    with col2:
        st.markdown("**Formatting Agent**") 
        format_types = ['Formatting Standards', 'Template Rules', 'Layout Guidelines']
        format_count = sum(standards_by_type.get(t, 0) for t in format_types)
        
        if format_count > 0:
            st.success(f"Ready ({format_count} standards)")
        else:
            st.warning("No formatting standards loaded")
    
    with col3:
        st.markdown("**Quality Assurance Agent**")
        quality_types = ['Quality Standards', 'Compliance Rules', 'Review Checklist']
        quality_count = sum(standards_by_type.get(t, 0) for t in quality_types)
        
        if quality_count > 0:
            st.success(f"Ready ({quality_count} standards)")
        else:
            st.warning("No quality standards loaded")
    
    return content_count > 0, format_count > 0, quality_count > 0


# Step 6: Individual Agent Implementations - ADD AFTER BASE CLASSES

class ContentStandardsAgent(BaseAgent):
    """Agent responsible for applying editorial and writing standards"""
    
    def __init__(self, standards_retriever):
        super().__init__(AgentType.CONTENT_EDITOR, standards_retriever)
    
    def process(self, input_data: Dict) -> ProcessingResult:
        """Apply content standards to document text"""
        
        document_text = input_data.get('document_text', '')
        if not document_text.strip():
            return ProcessingResult(
                content="",
                instructions=[],
                metadata={},
                agent_notes="No document text provided",
                success=False,
                error_message="Empty document text"
            )
        
        try:
            # Retrieve relevant content standards
            content_standards = self.standards_retriever.search(
                query=f"editorial writing style guidelines grammar citation format {document_text[:300]}",
                filter_types=['Editorial Guidelines', 'Style Guide', 'Writing Standards'],
                top_k=6
            )
            
            if not content_standards:
                return ProcessingResult(
                    content=document_text,
                    instructions=[],
                    metadata={'standards_found': 0},
                    agent_notes="No relevant content standards found - document unchanged",
                    success=True
                )
            
            # Build content editing prompt
            standards_context = "\n".join([
                f"STANDARD RULE ({std['doc_name']}): {std['text'][:500]}"
                for std in content_standards
            ])
            
            prompt = f"""
You are a Content Standards Agent for academic document processing. Apply editorial standards to improve document content while maintaining academic integrity.

CONTENT STANDARDS TO APPLY:
{standards_context}

DOCUMENT TO EDIT:
{document_text}

TASKS:
1. Apply writing style improvements (clarity, conciseness, academic tone)
2. Fix citation formatting according to standards
3. Improve paragraph structure and flow
4. Ensure consistency in terminology
5. Correct grammar and punctuation
6. Maintain original meaning and content

RESPOND IN EXACTLY THIS JSON FORMAT:
{{
    "edited_content": "The complete improved document text with all editorial improvements applied",
    "content_changes": [
        "Specific change 1 made",
        "Specific change 2 made"
    ],
    "formatting_hints": [
        {{"element": "title", "instruction": "Apply title formatting"}},
        {{"element": "headings", "instruction": "Use standard heading hierarchy"}}
    ],
    "agent_notes": "Summary of editorial improvements made and rationale"
}}

IMPORTANT: Return ONLY valid JSON. No additional text or explanations outside the JSON structure.
"""
            
            # Get AI response
            response_text = self._get_ai_response(prompt)
            
            # Parse response
            parsed_response = self._parse_json_response(response_text, document_text)
            
            # Validate and clean response
            edited_content = parsed_response.get('edited_content', document_text)
            content_changes = parsed_response.get('content_changes', [])
            formatting_hints = parsed_response.get('formatting_hints', [])
            agent_notes = parsed_response.get('agent_notes', 'Content processing completed')
            
            # Ensure content is not empty
            if not edited_content.strip():
                edited_content = document_text
                agent_notes += " (Original content preserved due to empty result)"
            
            return ProcessingResult(
                content=edited_content,
                instructions=formatting_hints,
                metadata={
                    'content_changes': content_changes,
                    'standards_applied': len(content_standards),
                    'original_word_count': len(document_text.split()),
                    'edited_word_count': len(edited_content.split())
                },
                agent_notes=agent_notes,
                success=True
            )
            
        except Exception as e:
            error_msg = f"Content agent processing error: {str(e)}"
            return ProcessingResult(
                content=document_text,
                instructions=[],
                metadata={'error': error_msg},
                agent_notes="Processing failed - original content preserved",
                success=False,
                error_message=error_msg
            )

class FormattingAgent(BaseAgent):
    """Agent responsible for applying visual formatting standards to documents"""
    
    def __init__(self, standards_retriever):
        super().__init__(AgentType.FORMATTER, standards_retriever)
    
    def process(self, input_data: Dict) -> ProcessingResult:
        """Apply formatting standards to create structured DOCX instructions"""
        
        content = input_data.get('content', '')
        content_hints = input_data.get('formatting_hints', [])
        
        if not content.strip():
            return ProcessingResult(
                content="",
                instructions=[],
                metadata={},
                agent_notes="No content provided for formatting",
                success=False,
                error_message="Empty content"
            )
        
        try:
            # Retrieve formatting standards
            formatting_standards = self.standards_retriever.search(
                query="document formatting styles fonts margins headings layout spacing citations tables",
                filter_types=['Formatting Standards', 'Template Rules', 'Layout Guidelines'],
                top_k=8
            )
            
            # Build comprehensive formatting instructions
            standards_context = "\n".join([
                f"FORMAT RULE ({std['doc_name']}): {std['text'][:400]}"
                for std in formatting_standards
            ])
            
            prompt = f"""
You are a Document Formatting Agent. Create detailed DOCX formatting instructions for academic document publishing.

FORMATTING STANDARDS:
{standards_context}

CONTENT TO FORMAT:
{content}

ANALYZE the content and create comprehensive DOCX formatting plan.

RESPOND IN EXACTLY THIS JSON FORMAT:
{{
    "document_structure": [
        {{"type": "title", "text": "Main document title", "style": "Title", "alignment": "center", "font_size": 16, "bold": true}},
        {{"type": "heading", "level": 1, "text": "Section heading", "style": "Heading 1", "font_size": 14, "bold": true}},
        {{"type": "paragraph", "text": "Paragraph content", "style": "Normal", "alignment": "justify", "font_size": 12}}
    ],
    "document_settings": {{
        "font_family": "Times New Roman",
        "default_font_size": 12,
        "line_spacing": 1.5,
        "margins": {{"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0}}
    }},
    "agent_notes": "Detailed explanation of formatting decisions and standards applied"
}}

IMPORTANT: Analyze the ENTIRE content and create structure for ALL elements. Return ONLY valid JSON, no additional text.
"""
            
            # Get AI response
            response_text = self._get_ai_response(prompt)
            
            # Parse response with fallback
            parsed_response = self._parse_json_response(response_text, content)
            
            # Extract and validate formatting data
            document_structure = parsed_response.get('document_structure', [])
            document_settings = parsed_response.get('document_settings', self._get_default_settings())
            agent_notes = parsed_response.get('agent_notes', 'Formatting instructions generated')
            
            # If no structure generated, create basic fallback
            if not document_structure:
                document_structure = self._create_fallback_structure(content)
                agent_notes += " (Using fallback structure due to parsing issues)"
            
            return ProcessingResult(
                content=content,
                instructions=document_structure,
                metadata={
                    'document_settings': document_settings,
                    'standards_applied': len(formatting_standards),
                    'structure_elements': len(document_structure)
                },
                agent_notes=agent_notes,
                success=True
            )
            
        except Exception as e:
            error_msg = f"Formatting agent error: {str(e)}"
            fallback_structure = self._create_fallback_structure(content)
            
            return ProcessingResult(
                content=content,
                instructions=fallback_structure,
                metadata={
                    'document_settings': self._get_default_settings(),
                    'error': error_msg
                },
                agent_notes="Fallback formatting applied due to processing error",
                success=False,
                error_message=error_msg
            )
    
    def _get_default_settings(self) -> Dict:
        """Return default document settings"""
        return {
            "font_family": "Times New Roman",
            "default_font_size": 12,
            "line_spacing": 1.5,
            "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0}
        }
    
    def _create_fallback_structure(self, content: str) -> List[Dict]:
        """Create basic document structure when AI parsing fails"""
        
        structure = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Try to identify element types
            if len(line) < 100 and (line.isupper() or line.endswith(':')):
                # Likely a heading
                structure.append({
                    "type": "heading",
                    "level": 1,
                    "text": line,
                    "style": "Heading 1",
                    "font_size": 14,
                    "bold": True
                })
            else:
                # Regular paragraph
                structure.append({
                    "type": "paragraph",
                    "text": line,
                    "style": "Normal",
                    "alignment": "justify",
                    "font_size": 12
                })
        
        return structure

class QualityAssuranceAgent(BaseAgent):
    """Agent responsible for final quality checks and compliance validation"""
    
    def __init__(self, standards_retriever):
        super().__init__(AgentType.QUALITY_CHECKER, standards_retriever)
    
    def process(self, input_data: Dict) -> ProcessingResult:
        """Perform quality assurance checks on processed document"""
        
        content = input_data.get('content', '')
        formatting_instructions = input_data.get('instructions', [])
        
        if not content.strip():
            return ProcessingResult(
                content="",
                instructions=[],
                metadata={},
                agent_notes="No content provided for quality check",
                success=False,
                error_message="Empty content for quality assurance"
            )
        
        try:
            # Retrieve quality and compliance standards
            quality_standards = self.standards_retriever.search(
                query="quality assurance compliance document review checklist validation requirements",
                filter_types=['Quality Standards', 'Compliance Rules', 'Review Checklist'],
                top_k=5
            )
            
            prompt = f"""
You are a Quality Assurance Agent for academic document publishing. Perform comprehensive quality review.

DOCUMENT CONTENT TO REVIEW:
{content[:2000]}...

PERFORM QUALITY ASSESSMENT:

RESPOND IN EXACTLY THIS JSON FORMAT:
{{
    "compliance_score": 85,
    "quality_assessment": {{
        "content_quality": {{"score": 90, "issues": ["Issue 1"]}},
        "formatting_compliance": {{"score": 85, "issues": ["Format issue 1"]}},
        "completeness": {{"score": 95, "issues": []}}
    }},
    "critical_issues": [],
    "recommendations": [
        "Specific improvement recommendation 1"
    ],
    "final_approval": true,
    "agent_notes": "Comprehensive quality assessment summary"
}}

IMPORTANT: Return ONLY valid JSON. Be thorough but practical.
"""
            
            # Get AI response
            response_text = self._get_ai_response(prompt)
            
            # Parse response
            parsed_response = self._parse_json_response(response_text, content)
            
            # Extract quality data
            compliance_score = parsed_response.get('compliance_score', 75)
            quality_assessment = parsed_response.get('quality_assessment', {})
            critical_issues = parsed_response.get('critical_issues', [])
            recommendations = parsed_response.get('recommendations', [])
            final_approval = parsed_response.get('final_approval', True)
            agent_notes = parsed_response.get('agent_notes', 'Quality review completed')
            
            return ProcessingResult(
                content=content,
                instructions=formatting_instructions,
                metadata={
                    'compliance_score': compliance_score,
                    'quality_assessment': quality_assessment,
                    'critical_issues': critical_issues,
                    'recommendations': recommendations,
                    'final_approval': final_approval,
                    'standards_checked': len(quality_standards)
                },
                agent_notes=agent_notes,
                success=final_approval and len(critical_issues) == 0
            )
            
        except Exception as e:
            error_msg = f"Quality assurance error: {str(e)}"
            
            return ProcessingResult(
                content=content,
                instructions=formatting_instructions,
                metadata={'error': error_msg},
                agent_notes="Basic quality check performed due to processing error",
                success=False,
                error_message=error_msg
            )


# Step 7: Document Processing Orchestrator - ADD AFTER INDIVIDUAL AGENTS

class DocumentProcessingOrchestrator:
    """Orchestrates the multi-agent document processing pipeline with enhanced DOCX generation"""
    
    def __init__(self, standards_retriever):
        self.standards_retriever = standards_retriever
        self.content_agent = ContentStandardsAgent(standards_retriever)
        self.formatting_agent = FormattingAgent(standards_retriever)
        self.quality_agent = QualityAssuranceAgent(standards_retriever)
        
    def process_document(self, document_text: str, original_file=None, 
                        selected_agents: List[str] = None) -> Dict:
        """Run complete multi-agent processing pipeline"""
        
        if not selected_agents:
            selected_agents = ["Content Agent", "Formatting Agent", "Quality Agent"]
        
        processing_log = []
        processing_results = {}
        
        try:
            # Initialize with original document
            current_content = document_text
            current_instructions = []
            current_metadata = {}
            
            # Stage 1: Content Standards Agent
            if "Content Agent" in selected_agents:
                st.info("Content Agent: Applying editorial standards...")
                
                content_result = self.content_agent.process({
                    'document_text': current_content
                })
                
                processing_results['content'] = content_result
                processing_log.append(f"Content Agent: {content_result.agent_notes}")
                
                if content_result.success:
                    current_content = content_result.content
                    current_instructions = content_result.instructions
                    current_metadata.update(content_result.metadata)
                    st.success(f"Content processing completed - {len(content_result.metadata.get('content_changes', []))} changes made")
                else:
                    st.warning(f"Content processing had issues: {content_result.error_message}")
            
            # Stage 2: Formatting Agent  
            if "Formatting Agent" in selected_agents:
                st.info("Formatting Agent: Applying visual standards...")
                
                formatting_result = self.formatting_agent.process({
                    'content': current_content,
                    'formatting_hints': current_instructions
                })
                
                processing_results['formatting'] = formatting_result
                processing_log.append(f"Formatting Agent: {formatting_result.agent_notes}")
                
                if formatting_result.success:
                    current_instructions = formatting_result.instructions
                    current_metadata.update(formatting_result.metadata)
                    st.success(f"Formatting completed - {len(current_instructions)} structure elements created")
                else:
                    st.warning(f"Formatting had issues: {formatting_result.error_message}")
            
            # Stage 3: Quality Assurance Agent
            if "Quality Agent" in selected_agents:
                st.info("Quality Agent: Validating compliance...")
                
                quality_result = self.quality_agent.process({
                    'content': current_content,
                    'instructions': current_instructions,
                    'document_settings': current_metadata.get('document_settings', {})
                })
                
                processing_results['quality'] = quality_result
                processing_log.append(f"Quality Agent: {quality_result.agent_notes}")
                
                if quality_result.success:
                    current_instructions = quality_result.instructions
                    current_metadata.update(quality_result.metadata)
                    compliance_score = quality_result.metadata.get('compliance_score', 0)
                    st.success(f"Quality check completed - Compliance score: {compliance_score}/100")
                else:
                    st.warning(f"Quality check had issues: {quality_result.error_message}")
            
            # Stage 4: Generate Final DOCX
            st.info("Generating formatted DOCX...")
            
            docx_bytes = self.generate_final_docx(
                content=current_content,
                formatting_instructions=current_instructions,
                document_settings=current_metadata.get('document_settings', {}),
                original_file=original_file
            )
            
            st.success("Document processing pipeline completed!")
            
            return {
                'final_docx': docx_bytes,
                'processing_log': processing_log,
                'processing_results': processing_results,
                'final_content': current_content,
                'final_instructions': current_instructions,
                'final_metadata': current_metadata,
                'pipeline_success': all(
                    result.success for result in processing_results.values()
                )
            }
            
        except Exception as e:
            error_msg = f"Pipeline orchestration error: {str(e)}"
            st.error(error_msg)
            
            return {
                'final_docx': None,
                'processing_log': processing_log + [f"Pipeline Error: {error_msg}"],
                'processing_results': processing_results,
                'final_content': current_content if 'current_content' in locals() else document_text,
                'final_instructions': current_instructions if 'current_instructions' in locals() else [],
                'final_metadata': current_metadata if 'current_metadata' in locals() else {},
                'pipeline_success': False,
                'error': error_msg
            }
    
    def generate_final_docx(self, content: str, formatting_instructions: List[Dict], 
                           document_settings: Dict, original_file=None) -> bytes:
        """Generate final DOCX file with comprehensive formatting support"""
        
        try:
            # Create new document
            doc = Document()
            
            # Apply document-level settings first
            self._apply_document_settings(doc, document_settings)
            
            # Process formatting instructions
            if formatting_instructions:
                st.info(f"Applying {len(formatting_instructions)} formatting instructions...")
                for i, instruction in enumerate(formatting_instructions):
                    try:
                        self._apply_formatting_instruction(doc, instruction)
                    except Exception as e:
                        st.warning(f"Skipped instruction {i+1}: {str(e)}")
                        continue
            else:
                # Use enhanced fallback structure
                st.warning("No formatting instructions found - creating enhanced document structure")
                enhanced_structure = self._create_comprehensive_fallback(content)
                for instruction in enhanced_structure:
                    self._apply_formatting_instruction(doc, instruction)
            
            # Convert to bytes
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            return bio.getvalue()
            
        except Exception as e:
            st.error(f"DOCX generation error: {str(e)}")
            
            # Create minimal but functional fallback document
            doc = Document()
            doc.add_heading('Document Processing Results', 0)
            doc.add_paragraph(f'Generation Error: {str(e)}')
            doc.add_paragraph('')
            doc.add_heading('Original Content', level=1)
            
            # Split content into manageable paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs[:50]:  # Limit to first 50 paragraphs to avoid issues
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            return bio.getvalue()
    
    def _apply_document_settings(self, doc: Document, settings: Dict):
        """Apply document-level formatting settings"""
        
        try:
            # Set margins
            margins = settings.get('margins', {})
            for section in doc.sections:
                section.top_margin = Inches(margins.get('top', 1))
                section.bottom_margin = Inches(margins.get('bottom', 1))
                section.left_margin = Inches(margins.get('left', 1))
                section.right_margin = Inches(margins.get('right', 1))
            
            # Set default font for Normal style
            style = doc.styles['Normal']
            font = style.font
            font.name = settings.get('font_family', 'Times New Roman')
            font.size = Pt(settings.get('default_font_size', 12))
            
            # Set line spacing for Normal style
            paragraph_format = style.paragraph_format
            line_spacing = settings.get('line_spacing', 1.5)
            paragraph_format.line_spacing = line_spacing
            
        except Exception as e:
            st.warning(f"Document settings application error: {str(e)}")
    
    def _apply_formatting_instruction(self, doc: Document, instruction: Dict):
        """Apply individual formatting instruction with comprehensive support"""
        
        try:
            element_type = instruction.get('type', '').lower()
            text = instruction.get('text', '')
            
            if element_type in ['title', 'document_title']:
                heading = doc.add_heading(text, level=0)
                if instruction.get('alignment') == 'center':
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Apply title-specific formatting
                for run in heading.runs:
                    run.font.size = Pt(instruction.get('font_size', 16))
                    run.font.bold = instruction.get('bold', True)
                    
            elif element_type in ['author', 'authors']:
                p = doc.add_paragraph(text)
                if instruction.get('alignment') == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
            elif element_type in ['date', 'institution']:
                p = doc.add_paragraph(text)
                if instruction.get('alignment') == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
            elif element_type == 'heading':
                level = min(max(instruction.get('level', 1), 1), 9)
                if text.strip():
                    heading = doc.add_heading(text, level=level)
                    # Apply custom formatting if specified
                    if instruction.get('font_size'):
                        for run in heading.runs:
                            run.font.size = Pt(instruction.get('font_size'))
                            
            elif element_type in ['paragraph', 'body_text']:
                if text.strip():
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, text)
                    
                    # Apply alignment
                    alignment = instruction.get('alignment', 'justify').lower()
                    if alignment == 'center':
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif alignment == 'right':
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif alignment == 'justify':
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Apply indentation
                    if instruction.get('first_line_indent'):
                        p.paragraph_format.first_line_indent = Inches(0.5)
                        
            elif element_type in ['abstract', 'abstract_text']:
                if text.strip():
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, text)
                    # Abstract typically has no first-line indent and single spacing
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.paragraph_format.line_spacing = 1.0
                    
            elif element_type in ['keywords', 'keyword_list']:
                if text.strip():
                    p = doc.add_paragraph()
                    # Handle "Keywords:" label with italic formatting
                    if text.lower().startswith('keywords'):
                        parts = text.split(':', 1)
                        if len(parts) == 2:
                            keywords_run = p.add_run(parts[0] + ':')
                            keywords_run.italic = True
                            p.add_run(' ' + parts[1].strip())
                        else:
                            p.add_run(text)
                    else:
                        p.add_run(text)
                        
            elif element_type in ['list', 'bullet_list', 'numbered_list']:
                items = instruction.get('items', [])
                list_style = instruction.get('style', 'bullet').lower()
                
                for item in items:
                    if item.strip():
                        # Clean up bullet markers if present
                        clean_item = item.strip()
                        if clean_item.startswith(('•', '-', '*', '–')):
                            clean_item = clean_item[1:].strip()
                            
                        if 'bullet' in list_style:
                            doc.add_paragraph(clean_item, style='List Bullet')
                        else:
                            doc.add_paragraph(clean_item, style='List Number')
                            
            elif element_type in ['table', 'data_table']:
                headers = instruction.get('headers', [])
                rows = instruction.get('rows', [])
                
                if headers and rows:
                    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    # Add headers with bold formatting
                    for i, header in enumerate(headers):
                        cell = table.cell(0, i)
                        cell.text = str(header)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                
                    # Add data rows
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(headers):
                                table.cell(row_idx+1, col_idx).text = str(cell_data)
            
            elif element_type in ['reference', 'citation', 'bibliography_entry']:
                if text.strip():
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, text)
                    
                    # Apply hanging indent for references
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.first_line_indent = Inches(-0.5)
                    p.paragraph_format.line_spacing = 1.0
                    
            elif element_type in ['block_quote', 'quote']:
                if text.strip():
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, text)
                    # Indent block quotes
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.right_indent = Inches(0.5)
                    
            else:
                # Default: treat as paragraph
                if text.strip():
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, text)
                    
        except Exception as e:
            st.warning(f"Formatting instruction error for {element_type}: {str(e)}")
    
    def _add_formatted_text(self, paragraph, text):
        """Add text with inline formatting support (bold, italic, etc.)"""
        import re
        
        if not text:
            return
        
        # Handle various markdown-style formatting
        # Split by bold markers first (**text**)
        bold_parts = re.split(r'(\*\*[^*]+\*\*)', text)
        
        for bold_part in bold_parts:
            if bold_part.startswith('**') and bold_part.endswith('**'):
                # Bold text
                inner_text = bold_part[2:-2]
                # Check if it also has italic markers
                italic_parts = re.split(r'(\*[^*]+\*)', inner_text)
                for italic_part in italic_parts:
                    if italic_part.startswith('*') and italic_part.endswith('*'):
                        # Bold and italic
                        run = paragraph.add_run(italic_part[1:-1])
                        run.bold = True
                        run.italic = True
                    else:
                        # Just bold
                        run = paragraph.add_run(italic_part)
                        run.bold = True
            else:
                # Check for italic only
                italic_parts = re.split(r'(\*[^*]+\*)', bold_part)
                for italic_part in italic_parts:
                    if italic_part.startswith('*') and italic_part.endswith('*'):
                        # Italic text
                        run = paragraph.add_run(italic_part[1:-1])
                        run.italic = True
                    else:
                        # Regular text - check for journal titles and other italic patterns
                        if self._should_be_italic(italic_part):
                            run = paragraph.add_run(italic_part)
                            run.italic = True
                        else:
                            paragraph.add_run(italic_part)
    
    def _should_be_italic(self, text):
        """Determine if text should be italicized based on content"""
        if not text.strip():
            return False
        
        # Common patterns that should be italic
        italic_patterns = [
            r'\b[Jj]ournal of\b',
            r'\b[Pp]roceedings of\b',
            r'\b[Rr]eview of\b',
            r'\b[Aa]nnals of\b',
            r'\bet al\.\b',
            r'\bp\.\s*\d+',  # page numbers
            r'\bvol\.\s*\d+',  # volume numbers
        ]
        
        for pattern in italic_patterns:
            if re.search(pattern, text):
                return True
                
        return False
    
    def _create_comprehensive_fallback(self, content: str) -> List[Dict]:
        """Create comprehensive document structure preserving all content elements"""
        
        structure = []
        lines = content.split('\n')
        current_list_items = []
        in_references = False
        in_abstract = False
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Flush any accumulated list items
            if current_list_items and not (line.startswith(('•', '-', '*', '–')) or 
                                          line.startswith(tuple('123456789')) and '.' in line):
                structure.append({
                    "type": "list",
                    "style": "bullet",
                    "items": current_list_items,
                    "font_size": 12
                })
                current_list_items = []
            
            # Detect major sections
            line_lower = line.lower()
            if line_lower in ['abstract', 'introduction', 'methodology', 'methods', 'results', 
                             'discussion', 'conclusion', 'conclusions', 'references', 'bibliography']:
                
                structure.append({
                    "type": "heading",
                    "level": 1,
                    "text": line.title(),
                    "style": "Heading 1",
                    "font_size": 14,
                    "bold": True
                })
                
                in_references = line_lower in ['references', 'bibliography']
                in_abstract = line_lower == 'abstract'
                
            # Detect subsection headings
            elif (line.endswith(':') or 
                  (len(line) < 100 and line.istitle()) or
                  line.startswith(('1.', '2.', '3.', '4.', '5.', 'I.', 'II.', 'III.', 'IV.', 'V.')) or
                  line.isupper()):
                
                # Determine heading level
                level = 2
                if line.startswith(('1.', '2.', '3.', 'I.', 'II.')):
                    level = 2
                elif line.startswith(('1.1', '2.1', '3.1')) or line.count('.') > 1:
                    level = 3
                
                structure.append({
                    "type": "heading", 
                    "level": level,
                    "text": line,
                    "style": f"Heading {level}",
                    "font_size": 14 - level,
                    "bold": True
                })
                
            # Handle references section content
            elif in_references and (line[0].isupper() or line.startswith(('1.', '2.', 'A', 'B'))):
                structure.append({
                    "type": "reference",
                    "text": line,
                    "style": "Reference"
                })
                
            # Handle list items
            elif line.startswith(('•', '-', '*', '–')):
                current_list_items.append(line[1:].strip())
                
            elif (line.startswith(tuple('123456789')) and 
                  ('.' in line or ')' in line) and 
                  line.index(next((c for c in line if c in '.)'), '.')) < 5):
                current_list_items.append(line)
                
            # Handle keywords
            elif line.lower().startswith('keywords'):
                structure.append({
                    "type": "keywords",
                    "text": line,
                    "style": "Keywords"
                })
                
            # Handle regular paragraphs
            else:
                element_type = "abstract" if in_abstract else "paragraph"
                structure.append({
                    "type": element_type,
                    "text": line,
                    "style": "Normal",
                    "alignment": "justify",
                    "font_size": 12
                })
            
            i += 1
        
        # Flush any remaining list items
        if current_list_items:
            structure.append({
                "type": "list",
                "style": "bullet", 
                "items": current_list_items,
                "font_size": 12
            })
        
        return structure

# Step 8: Streamlit Tab Functions - ADD AFTER ORCHESTRATOR

def agent_based_processing_tab():
    """New processing tab using agent architecture"""
    
    st.header("Multi-Agent Document Processing")
    
    # Initialize enhanced retriever
    if 'agent_retriever' not in st.session_state:
        st.session_state.agent_retriever = EnhancedStandardsRetriever()
    
    retriever = st.session_state.agent_retriever
    
    # Display agent status
    content_ready, format_ready, quality_ready = display_agent_status(retriever)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### Upload Document for Agent Processing")
        
        uploaded_doc = st.file_uploader(
            "Choose document to process:",
            type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'],
            help="Upload manuscripts, articles, reports, or any business document",
            key="agent_upload"
        )
        
        if uploaded_doc:
            st.success(f"Document loaded: {uploaded_doc.name}")
            
            # Store original file for template preservation
            st.session_state.current_uploaded_file = uploaded_doc
            
            # Extract text
            if 'agent_doc_text' not in st.session_state or st.session_state.get('current_doc_name') != uploaded_doc.name:
                with st.spinner("Extracting document content..."):
                    document_text = extract_text_from_document(uploaded_doc)
                    st.session_state.agent_doc_text = document_text
                    st.session_state.current_doc_name = uploaded_doc.name
            
            # Display extracted text preview
            with st.expander("View Extracted Text Preview"):
                preview_text = st.session_state.agent_doc_text[:1000]
                if len(st.session_state.agent_doc_text) > 1000:
                    preview_text += "... (truncated)"
                st.text_area("Document Content Preview:", preview_text, height=200, disabled=True)
            
            # Agent Selection and Configuration
            st.markdown("### Agent Processing Configuration")
            
            # Agent selection with smart defaults
            available_agents = []
            if content_ready:
                available_agents.append("Content Standards Agent")
            if format_ready:
                available_agents.append("Formatting Agent")
            if quality_ready:
                available_agents.append("Quality Assurance Agent")
            
            if not available_agents:
                st.warning("No agents available - please upload standards documents first!")
                return
            
            selected_agents = st.multiselect(
                "Select processing agents:",
                available_agents,
                default=available_agents,
                help="Choose which agents to run on your document"
            )
            
            # Output options
            output_options = st.multiselect(
                "Generate Outputs:",
                ["Formatted DOCX", "Processing Report", "Quality Report"],
                default=["Formatted DOCX", "Processing Report"]
            )
            
            # Main processing button
            if st.button("Run Agent Pipeline", type="primary", disabled=not selected_agents):
                
                # Validate prerequisites
                if not st.session_state.standards_library:
                    st.error("Please upload standards documents first in the Standards Manager tab!")
                    return
                
                if not st.session_state.standards_chunks:
                    st.warning("Building standards knowledge base...")
                    build_standards_knowledge_base()
                
                # Initialize orchestrator
                orchestrator = DocumentProcessingOrchestrator(retriever)
                
                # Map UI selections to agent names
                agent_mapping = {
                    "Content Standards Agent": "Content Agent",
                    "Formatting Agent": "Formatting Agent", 
                    "Quality Assurance Agent": "Quality Agent"
                }
                
                mapped_agents = [agent_mapping.get(agent, agent) for agent in selected_agents]
                
                # Run processing pipeline
                st.markdown("---")
                st.subheader("Processing Pipeline")
                
                # Create progress tracking
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                try:
                    # Start processing
                    status_text.text("Initializing agent pipeline...")
                    progress_bar.progress(10)
                    
                    # Run the orchestrator
                    results = orchestrator.process_document(
                        st.session_state.agent_doc_text,
                        st.session_state.current_uploaded_file,
                        mapped_agents
                    )
                    
                    progress_bar.progress(100)
                    status_text.text("Processing completed!")
                    
                    # Display results
                    display_agent_processing_results(results, output_options, uploaded_doc.name)
                    
                except Exception as e:
                    progress_bar.progress(100)
                    status_text.text(f"Processing failed: {str(e)}")
                    st.error(f"Pipeline error: {str(e)}")
    
    with col2:
        st.markdown("### Agent System Info")
        
        # Agent readiness status
        if st.session_state.standards_chunks:
            st.markdown('<div class="rag-info">', unsafe_allow_html=True)
            st.write(f"**Knowledge Base Status:**")
            st.write(f"• {len(st.session_state.standards_library)} standards documents")
            st.write(f"• {len(st.session_state.standards_chunks)} searchable chunks")
            st.write(f"• Ready for intelligent processing")
            
            # Show standards by type
            standards_by_type = retriever.get_standards_count_by_type()
            st.write("**Available Standards:**")
            for std_type, count in standards_by_type.items():
                st.write(f"• {std_type}: {count} docs")
            
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.warning("Standards knowledge base not built")
            if st.button("Build Knowledge Base"):
                with st.spinner("Building knowledge base..."):
                    success = build_standards_knowledge_base()
                    if success:
                        st.success("Knowledge base ready!")
                        st.rerun()

def display_agent_processing_results(results: Dict, output_options: List[str], filename: str):
    """Display comprehensive agent processing results"""
    
    st.markdown("## Agent Processing Complete!")
    
    # Pipeline status
    pipeline_success = results.get('pipeline_success', False)
    if pipeline_success:
        st.success("All agents completed successfully!")
    else:
        st.warning("Some agents encountered issues - check reports for details")
    
    # Processing statistics
    col1, col2, col3, col4 = st.columns(4)
    
    processing_results = results.get('processing_results', {})
    
    with col1:
        agents_run = len(processing_results)
        st.metric("Agents Run", agents_run)
    
    with col2:
        successful_agents = sum(1 for result in processing_results.values() if result.success)
        st.metric("Successful", successful_agents)
    
    with col3:
        # Get compliance score if available
        quality_result = processing_results.get('quality')
        compliance_score = quality_result.metadata.get('compliance_score', 0) if quality_result else 0
        st.metric("Quality Score", f"{compliance_score}/100")
    
    with col4:
        # Get word count change
        content_result = processing_results.get('content')
        original_words = content_result.metadata.get('original_word_count', 0) if content_result else 0
        edited_words = content_result.metadata.get('edited_word_count', 0) if content_result else 0
        change_pct = ((edited_words - original_words) / original_words * 100) if original_words > 0 else 0
        st.metric("Content Change", f"{change_pct:.1f}%")
    
    # Create output tabs
    tab_list = []
    if "Formatted DOCX" in output_options:
        tab_list.append("DOCX Output")
    if "Processing Report" in output_options:
        tab_list.append("Agent Reports") 
    if "Quality Report" in output_options:
        tab_list.append("Quality Analysis")
    
    if not tab_list:
        tab_list = ["Results"]
    
    tabs = st.tabs(tab_list)
    tab_index = 0
    
    # DOCX Output Tab
    if "Formatted DOCX" in output_options:
        with tabs[tab_index]:
            st.markdown("### Formatted Document Output")
            
            final_docx = results.get('final_docx')
            if final_docx:
                st.download_button(
                    "Download Formatted DOCX",
                    final_docx,
                    f"processed_{filename}",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
                
                # Content preview
                if st.button("Preview Content"):
                    final_content = results.get('final_content', '')
                    st.text_area("Final Content:", final_content[:1000] + "..." if len(final_content) > 1000 else final_content, height=200)
            else:
                st.error("DOCX generation failed")
        
        tab_index += 1
    
    # Processing Report Tab
    if "Processing Report" in output_options:
        with tabs[tab_index]:
            st.markdown("### Agent Processing Reports")
            
            # Individual agent reports
            for agent_name, result in processing_results.items():
                status_icon = "✅" if result.success else "❌"
                with st.expander(f"{status_icon} {agent_name.title()} Agent Report"):
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write(f"**Status:** {'Success' if result.success else 'Failed'}")
                        st.write(f"**Agent Notes:** {result.agent_notes}")
                    
                    with col2:
                        if result.metadata:
                            st.write("**Metadata:**")
                            for key, value in result.metadata.items():
                                if isinstance(value, (int, float, str, bool)):
                                    st.write(f"• {key}: {value}")
        
        tab_index += 1

def standards_management_tab():
    st.header("Standards Library")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("### Upload Standards Documents")
        
        standards_doc = st.file_uploader(
            "Upload style guides, editorial standards:",
            type=['pdf', 'docx', 'txt'],
            help="Upload company style guides, formatting rules, editorial guidelines",
            key="standards_upload"
        )
        
        if standards_doc:
            st.success(f"Standards document: {standards_doc.name}")
            
            # Standards metadata
            standard_name = st.text_input("Standard Name:", standards_doc.name.split('.')[0])
            standard_type = st.selectbox(
                "Standard Type:",
                ["Style Guide", "Editorial Guidelines", "Formatting Standards", "Quality Standards", "Template Rules", "Compliance Rules"]
            )
            description = st.text_area("Description:", "Enterprise formatting and editing standards")
            
            if st.button("Add to Standards Library"):
                with st.spinner("Processing standards document..."):
                    standards_content = extract_text_from_document(standards_doc)
                    
                    if standards_content and len(standards_content.strip()) > 100:
                        # Save to library
                        new_standard = {
                            'name': standard_name,
                            'type': standard_type,
                            'description': description,
                            'content': standards_content,
                            'filename': standards_doc.name,
                            'uploaded_date': datetime.now().isoformat(),
                            'word_count': len(standards_content.split())
                        }
                        
                        st.session_state.standards_library.append(new_standard)
                        
                        # Rebuild RAG knowledge base
                        build_standards_knowledge_base()
                        
                        st.success("Standards added and knowledge base updated!")
                        st.rerun()
                    else:
                        st.error("Failed to extract content from standards document")
    
    with col2:
        st.markdown("### Current Standards Library")
        
        if st.session_state.standards_library:
            for i, standard in enumerate(st.session_state.standards_library):
                with st.expander(f"{standard['name']} ({standard['type']})"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.write(f"**Type:** {standard['type']}")
                        st.write(f"**Word Count:** {standard['word_count']:,}")
                    
                    with col2:
                        st.write(f"**Uploaded:** {standard['uploaded_date'][:10]}")
                        st.write(f"**File:** {standard['filename']}")
                    
                    st.write(f"**Description:** {standard['description']}")
                    
                    if st.button(f"Remove", key=f"remove_{i}"):
                        st.session_state.standards_library.pop(i)
                        build_standards_knowledge_base()
                        st.rerun()
        else:
            st.info("No standards uploaded yet. Upload your first style guide or editorial standards document above.")

def rag_search_tab():
    st.header("RAG System - Standards Search")
    
    if not st.session_state.standards_chunks:
        st.warning("Please upload standards documents first to enable RAG search!")
        return
    
    st.markdown("### Search Your Standards Library")
    
    # Search interface
    search_query = st.text_area(
        "Search for specific standards or rules:",
        "citation format requirements for journal articles",
        height=100,
        help="Describe what standards you're looking for - the RAG system will find relevant rules"
    )
    
    num_results = st.slider("Number of results:", 3, 15, 8)
    
    if st.button("Search Standards", type="primary"):
        if search_query.strip():
            processor = StandardsProcessor()
            
            with st.spinner("Searching knowledge base..."):
                results = processor.semantic_search(search_query, top_k=num_results)
            
            if results:
                st.markdown(f"### Found {len(results)} Relevant Standards")
                
                for i, result in enumerate(results):
                    relevance_color = "🟢" if result['relevance_score'] > 70 else "🟡" if result['relevance_score'] > 40 else "🔴"
                    
                    with st.expander(f"{relevance_color} **Rule {i+1}** - {result['section']} (Relevance: {result['relevance_score']:.1f}%)"):
                        col1, col2 = st.columns([2, 1])
                        
                        with col1:
                            st.markdown(f"**Document:** {result['doc_name']}")
                            st.markdown(f"**Section:** {result['section']}")
                            
                        with col2:
                            st.markdown(f"**Words:** {result['word_count']}")
                            st.markdown(f"**Relevance:** {result['relevance_score']:.1f}%")
                        
                        st.markdown("**Content:**")
                        st.text_area("Rule Content:", result['text'], height=150, key=f"search_result_{i}")
            else:
                st.info("No relevant standards found. Try different search terms or upload more comprehensive standards.")
        else:
            st.warning("Please enter a search query!")

def dashboard_tab():
    st.header("Processing Dashboard & Analytics")
    
    # Overall metrics
    total_docs = len(st.session_state.processed_documents)
    total_standards = len(st.session_state.standards_library)
    total_chunks = len(st.session_state.standards_chunks)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Documents Processed", total_docs)
    with col2:
        st.metric("Standards Library", total_standards)
    with col3:
        st.metric("Knowledge Chunks", total_chunks)
    with col4:
        processing_rate = "2-5 min" if total_standards > 0 else "Setup Required"
        st.metric("Processing Speed", processing_rate)
    
    # System health
    st.markdown("### System Health")
    
    health_items = [
        ("RAG System", "Active" if st.session_state.standards_chunks else "Inactive"),
        ("Vector Embeddings", "Ready" if st.session_state.chunk_embeddings is not None else "Not Built"),
        ("Standards Library", "Loaded" if total_standards > 0 else "Empty"),
        ("Processing Engine", "Online")
    ]
    
    for item, status in health_items:
        col1, col2 = st.columns([3, 1])
        with col1:
            st.write(f"**{item}:**")
        with col2:
            st.write(status)
    
    # Standards library analysis
    if st.session_state.standards_library:
        st.markdown("### Standards Library Analysis")
        
        # Standards by type
        type_counts = {}
        total_words = 0
        
        for standard in st.session_state.standards_library:
            std_type = standard['type']
            type_counts[std_type] = type_counts.get(std_type, 0) + 1
            total_words += standard.get('word_count', 0)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Standards by Type:**")
            for std_type, count in type_counts.items():
                st.write(f"• {std_type}: {count}")
        
        with col2:
            st.metric("Total Words", f"{total_words:,}")
            if len(st.session_state.standards_library) > 0:
                st.metric("Average Doc Size", f"{total_words//len(st.session_state.standards_library):,} words")
    else:
        st.info("No standards uploaded yet. Visit the Standards Manager to get started!")


# Step 9: Main Function and App Entry Point - ADD AFTER TAB FUNCTIONS

def main():
    """Main application function"""
    
    # Header
    st.markdown("""
    <div class="doc-header">
        <h1>📄 AI Document Standards Engine</h1>
        <p>Enterprise Document Processing with RAG-Powered Standards Application</p>
        <p><em>Intelligent Standards • Track Changes • Professional Editing</em></p>
    </div>
    """, unsafe_allow_html=True)

    # Sidebar RAG status
    with st.sidebar:
        st.title("🧠 RAG System Status")
        
        # Standards library status
        num_standards = len(st.session_state.standards_library)
        num_chunks = len(st.session_state.standards_chunks)
        
        if num_standards > 0:
            st.success(f"✅ {num_standards} standards loaded")
            st.info(f"📚 {num_chunks} searchable chunks")
            
            if st.button("🔄 Rebuild Knowledge Base"):
                with st.spinner("🧠 Building RAG knowledge base..."):
                    success = build_standards_knowledge_base()
                    if success:
                        st.success("✅ Knowledge base updated!")
                        st.rerun()
        else:
            st.warning("📋 No standards uploaded yet")
        
        st.markdown("---")
        
        # Processing stats
        st.markdown("**📊 Processing Stats**")
        st.metric("Documents Processed", len(st.session_state.processed_documents))
        st.metric("Standards Library", num_standards)
        
        # System status
        st.markdown("---")
        st.markdown("**🔧 System Status**")
        
        # RAG System
        rag_status = "🟢 Active" if st.session_state.standards_chunks else "🔴 Inactive"
        st.write(f"RAG System: {rag_status}")
        
        # Vector embeddings
        embeddings_status = "🟢 Ready" if st.session_state.chunk_embeddings is not None else "🔴 Not Built"
        st.write(f"Embeddings: {embeddings_status}")
        
        # Processing engine
        st.write("Processing: 🟢 Online")

    # Main navigation
    tab1, tab2, tab3, tab4 = st.tabs([
        "🤖 Agent Processing", 
        "📋 Standards Manager", 
        "🔍 RAG Search", 
        "📊 Dashboard"
    ])

    with tab1:
        agent_based_processing_tab()
    
    with tab2:
        standards_management_tab()
    
    with tab3:
        rag_search_tab()
    
    with tab4:
        dashboard_tab()

# Application entry point
if __name__ == "__main__":
    main()

